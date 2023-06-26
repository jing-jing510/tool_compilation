import tkinter as tk
from tkinter import filedialog
from docx import Document
import re

def combine_lines(doc):
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip() != '']
    new_paragraphs = []
    temp = ''

    for p in paragraphs:
        if re.match(r"\d+\.", p.strip()):
            if temp:
                new_paragraphs.append(temp.strip())
                temp = ''
        temp += ' ' + p.strip()

    if temp:
        new_paragraphs.append(temp.strip())
    return new_paragraphs

def process_files():
    # 获取用户输入的文件名
    input_file = filedialog.askopenfilename(filetypes=[("Word files", "*.docx")])
    output_file = filedialog.asksaveasfilename(defaultextension=".docx")

    # 读取Word文件
    doc = Document(input_file)

    # 修改内容
    new_paragraphs = combine_lines(doc)

    # 创建一个新的文档
    new_doc = Document()

    # 将修改后的内容添加到新的文档中
    for p in new_paragraphs:
        new_doc.add_paragraph(p)  # 添加一个空段落作为额外的换行

    # 保存到新的文件中
    new_doc.save(output_file)

root = tk.Tk()
root.withdraw()  # 隐藏主窗口

process_files()

root.mainloop()
